import os
import re
from collections import Counter


def parse_file(path):
    """Returns (words, full_text, word_offsets, page_starts).
    page_starts is a list of word indices (one per page with content)
    for PDFs, or None for other formats.
    """
    ext = os.path.splitext(path)[1].lower()

    if ext == ".pdf":
        return _parse_pdf_pages(path)

    parsers = {
        ".txt": parse_txt,
        ".docx": parse_docx,
        ".epub": parse_epub,
        ".html": parse_html,
        ".htm": parse_html,
        ".rtf": parse_rtf,
        ".odt": parse_odt,
        ".md": parse_txt,
        ".csv": parse_txt,
        ".xml": parse_txt,
        ".json": parse_txt,
    }
    parser = parsers.get(ext, parse_txt_raw)
    raw = parser(path)
    text = clean_text(raw, from_pdf=False)
    words, offsets = tokenize_with_offsets(text)
    return words, text, offsets, None


def _parse_pdf_pages(path):
    """Parse PDF page by page, filtering headers/footers via Y-zone +
    content-frequency detection. Returns (words, full_text, offsets, page_starts)."""
    import fitz

    doc = fitz.open(path)
    try:
        chrome_prefixes = _detect_page_chrome(doc)

        per_page_cleaned = []
        for page in doc:
            ph = page.rect.height
            blocks = page.get_text("dict")["blocks"]
            lines = []
            for b in blocks:
                if b.get("type") != 0:
                    continue
                y0, y1 = b["bbox"][1], b["bbox"][3]
                txt = " ".join(
                    s["text"] for ld in b["lines"] for s in ld["spans"]
                ).strip()
                if not txt:
                    continue
                # Extreme top (running heads, page numbers): always filter
                if y1 < 0.08 * ph:
                    continue
                # Extreme bottom (copyright, page numbers): always filter
                if y0 > 0.92 * ph:
                    continue
                # Header / footer margin zone: filter if content matches chrome
                if (y1 < 0.12 * ph or y0 > 0.85 * ph) and _matches_chrome(txt, chrome_prefixes):
                    continue
                for ld in b["lines"]:
                    line = "".join(s["text"] for s in ld["spans"])
                    lines.append(line)

            if not lines:
                raw = page.get_text()
                if raw.strip():
                    lines = raw.split("\n")
            if lines:
                cleaned = clean_text("\n".join(lines), from_pdf=True)
                if cleaned:
                    per_page_cleaned.append(cleaned)

        if not per_page_cleaned and chrome_prefixes:
            # Chrome detection was too aggressive for this document.
            # Fall back to plain extraction with no Y-zone filtering.
            for page in doc:
                raw = page.get_text()
                if raw.strip():
                    cleaned = clean_text(raw, from_pdf=True)
                    if cleaned:
                        per_page_cleaned.append(cleaned)

        full_text = "\n\n".join(per_page_cleaned)
        words, offsets = tokenize_with_offsets(full_text)

        page_starts = []
        word_idx = 0
        for cleaned in per_page_cleaned:
            count = len(re.findall(r"\S+", cleaned))
            if count > 0:
                page_starts.append(word_idx)
                word_idx += count

        return words, full_text, offsets, page_starts
    finally:
        doc.close()


def _detect_page_chrome(doc, sample_limit=15):
    """Sample pages to build a set of header/footer 'chrome' text prefixes.

    Collects blocks from the header zone (y1 < 12% page height) and footer
    zone (y0 > 90% page height). Page-number suffixes (\" | 15\") are
    normalised before counting so variants of the same running header are
    detected as one pattern.

    Returns a set of 60-char prefixes used to filter repeated header/footer
    text during extraction.
    """
    n = len(doc)
    if n < 4:
        return set()

    start = max(0, n // 7)
    sample_pages = list(range(start, min(start + sample_limit, n)))

    header_texts = []
    footer_texts = []

    for pno in sample_pages:
        page = doc[pno]
        ph = page.rect.height
        for b in page.get_text("dict")["blocks"]:
            if b.get("type") != 0:
                continue
            y0, y1 = b["bbox"][1], b["bbox"][3]
            txt = " ".join(
                s["text"] for ld in b["lines"] for s in ld["spans"]
            ).strip()
            if not txt:
                continue
            norm = _norm_chrome_text(txt)
            if y1 < 0.12 * ph:
                header_texts.append(norm)
            elif y0 > 0.90 * ph:
                footer_texts.append(norm)

    h_freq = Counter(header_texts)
    f_freq = Counter(footer_texts)
    threshold = max(2, len(sample_pages) // 5)

    chrome_prefixes = set()
    for freq in (h_freq, f_freq):
        for text, count in freq.items():
            if count >= threshold and len(text) >= 30:
                chrome_prefixes.add(text[:60])

    return chrome_prefixes


def _norm_chrome_text(txt):
    """Normalise header/footer text by removing trailing page-number patterns
    so 'Los Libros Mánticos de Suroeste | 15' and '... | 16' compare equal."""
    return re.sub(r"\s*[|\-–—]\s*\d+\s*$", "", txt).strip()


def _matches_chrome(text, chrome_prefixes):
    """Check if text matches a known chrome prefix (exact or with page-number suffix)."""
    if not chrome_prefixes:
        return False
    for prefix in chrome_prefixes:
        if text[: len(prefix)] == prefix:
            return True
        rest = text[len(prefix):]
        if re.match(r"^\s*[|\-–—]\s*\d+\s*$", rest):
            return True
    return False
    for prefix in chrome_prefixes:
        if text[: len(prefix)] == prefix:
            return True
        # Variant: prefix + " | N" trailing page number or "14 |" prefix
        rest = text[len(prefix):]
        if re.match(r"^\s*[|\-–—]\s*\d+\s*$", rest):
            return True
    return False


_SPACED_LETTERS_RE = re.compile(
    r"(?<![\w])((?:[A-Za-záéíóúüñÁÉÍÓÚÜÑ] ){3,}[A-Za-záéíóúüñÁÉÍÓÚÜÑ])(?![\w])"
)


def _collapse_spaced_letters(text):
    """Join runs of >=4 single-letter tokens: 'a b s t r a c t' -> 'abstract'.
    Runs are space-delimited only (newlines break runs), so multi-line
    sidebar headers collapse per line into separate words."""
    return _SPACED_LETTERS_RE.sub(lambda m: m.group(1).replace(" ", ""), text)


def clean_text(text, from_pdf=False):
    """Clean extracted text for RSVP.
    from_pdf=True enables PDF-specific rules (page numbers,
    TOC entries, column reorder, missing-space insertion).
    Other formats use only universal rules.
    """
    # 1. Join hyphenated words broken across lines
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)

    # 1.5. Collapse letter-spaced runs into real words.
    # PDF sidebars (Elsevier "a r t i c l e  i n f o") and letterspaced
    # headers produce runs of single-letter tokens that pollute RSVP.
    # Threshold >= 4 avoids false positives like "A y B".
    text = _collapse_spaced_letters(text)

    lines = text.split("\n")

    if from_pdf:
        # 2. Remove isolated page numbers
        lines = [l for l in lines if not re.match(r"^\s*\d+\s*$", l)]

        # 3. Remove common TOC entries: "text ...... 123" or "text   123"
        cleaned = []
        for l in lines:
            stripped = l.strip()
            if not stripped:
                cleaned.append("")
                continue
            if re.search(r"\s[\.\s]{8,}\s\d+$", stripped):
                continue
            cleaned.append(l)
        lines = cleaned

        # 4. Detect and reorder multi-column layout (per contiguous block)
        lines = _reorder_columns(lines)

        # 5. Fix missing spaces from PDF extraction
        text_lines = "\n".join(lines)
        text_lines = re.sub(
            r"(?<=[a-záéíóúüñ])(?=[A-ZÁÉÍÓÚÜÑ])", " ", text_lines
        )
        text_lines = re.sub(
            r"(?<=\d)(?=[A-Za-záéíóúüñÁÉÍÓÚÜÑ])", " ", text_lines
        )
        lines = text_lines.split("\n")

    text = "\n".join(lines)

    # 6. Split words at double-hyphen/em-dash/en-dash for RSVP readability
    # (U+2212 minus is intentionally NOT split: it is a math operator)
    text = re.sub(r"--(\S)", r" --\1", text)
    text = re.sub(r"—(\S)", r" —\1", text)
    text = re.sub(r"–(\S)", r" –\1", text)

    # 7. Normalize whitespace (preserve paragraph breaks)
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)

    return text.strip()


def _reorder_columns(lines):
    """Detect and reorder two-column text within contiguous blocks.
    Interleaves left/right lines for natural reading order."""
    result = []
    col_left = []
    col_right = []
    gap_re = re.compile(r"  {4,}")

    for l in lines:
        m = gap_re.search(l)
        if m and len(l) > 40:
            mid = m.start() + (m.end() - m.start()) // 2
            left = l[:mid].strip()
            right = l[mid:].strip()
            if left and right and len(left) > 10 and len(right) > 10:
                col_left.append(left)
                col_right.append(right)
                continue
        # Not a column line: flush any pending column block
        if col_left or col_right:
            result.extend(_interleave_columns(col_left, col_right))
            col_left = []
            col_right = []
        result.append(l)

    # Flush remaining columns
    if col_left or col_right:
        result.extend(_interleave_columns(col_left, col_right))

    return result


def _interleave_columns(left, right):
    """Interleave left and right column lines for natural reading flow."""
    interleaved = []
    max_len = max(len(left), len(right))
    for i in range(max_len):
        if i < len(left):
            interleaved.append(left[i])
        if i < len(right):
            interleaved.append(right[i])
    return interleaved


def tokenize_with_offsets(text):
    words = []
    offsets = []
    for m in re.finditer(r"\S+", text):
        words.append(m.group())
        offsets.append([m.start(), m.end()])
    return words, offsets


def parse_txt(path):
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def parse_txt_raw(path):
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            raw = f.read()
        if len(raw) > 100:
            return raw
        with open(path, "rb") as f:
            raw = f.read().decode("utf-8", errors="replace")
        return raw
    except Exception:
        return ""


def parse_docx(path):
    from docx import Document

    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    cells.append(cell_text)
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def parse_epub(path):
    import ebooklib
    from ebooklib import epub
    from bs4 import BeautifulSoup

    book = epub.read_epub(path)
    texts = []

    # Use spine order for correct reading sequence
    try:
        for doc_id, _ in book.spine:
            item = book.get_item_with_id(doc_id)
            if item and item.get_type() == ebooklib.ITEM_DOCUMENT:
                soup = BeautifulSoup(item.get_content(), "html.parser")
                texts.append(soup.get_text())
    except Exception:
        texts = []

    # Fallback if spine produced nothing
    if not texts:
        try:
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    soup = BeautifulSoup(item.get_content(), "html.parser")
                    texts.append(soup.get_text())
        except Exception:
            pass

    return "\n".join(texts)


def parse_html(path):
    from bs4 import BeautifulSoup

    with open(path, "r", encoding="utf-8", errors="replace") as f:
        soup = BeautifulSoup(f.read(), "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
            tag.decompose()
        return soup.get_text(separator="\n")


def parse_rtf(path):
    from striprtf.striprtf import rtf_to_text

    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return rtf_to_text(f.read())


def parse_odt(path):
    import zipfile
    from xml.etree import ElementTree

    NS = "urn:oasis:names:tc:opendocument:xmlns:text:1.0"
    with zipfile.ZipFile(path) as z:
        content = z.read("content.xml")
        root = ElementTree.fromstring(content)
        texts = []
        for node in root.iter():
            tag = node.tag
            if tag in (f"{{{NS}}}p", f"{{{NS}}}h"):
                texts.append("".join(node.itertext()))
        return "\n".join(texts)
